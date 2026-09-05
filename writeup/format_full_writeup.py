"""Apply a plain research-document layout to the pandoc-generated DOCX.

Usage: python format_full_writeup.py INPUT.docx OUTPUT.docx
The input document is left unchanged.
"""

import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

d = Document(sys.argv[1])
for section in d.sections:
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    section.top_margin = section.bottom_margin = Inches(0.8)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.35)
for name in ["Normal", "Body Text", "First Paragraph", "Compact", "Block Text", "Caption"]:
    s = d.styles[name]
    s.font.name = "Arial"
    s.font.size = Pt(11 if name not in ["Block Text", "Caption"] else 10)
    s.font.color.rgb = RGBColor(0, 0, 0)
    s.paragraph_format.line_spacing = 1.05
    s.paragraph_format.space_after = Pt(6)
    s.paragraph_format.widow_control = True
    s.paragraph_format.keep_together = False
for name, size in [("Heading 1", 17), ("Heading 2", 14), ("Heading 3", 12), ("Heading 4", 11)]:
    s = d.styles[name]
    s.font.name, s.font.size = "Arial", Pt(size)
    s.font.bold = True
    s.font.color.rgb = RGBColor(0, 0, 0)
    s.paragraph_format.space_before = Pt(12)
    s.paragraph_format.space_after = Pt(6)
    s.paragraph_format.keep_with_next = True
    s.paragraph_format.keep_together = True
quote_style = d.styles["Block Text"]
quote_style.paragraph_format.left_indent = Inches(0.18)
quote_style.paragraph_format.right_indent = Inches(0.05)
quote_style.paragraph_format.space_after = Pt(4)
for p in d.paragraphs:
    if p.text == "A randomly selected endpoint example":
        p.paragraph_format.page_break_before = True
    if p._p.xpath(".//w:drawing"):
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.keep_together = True
    if p.text.startswith("Figure ") or p.style.name == "Caption":
        p.paragraph_format.keep_together = True
        p.paragraph_format.space_after = Pt(6)
    if p.text.startswith("Safety drift under explicit"):
        p.style = d.styles["Normal"]
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.keep_with_next = True
        for r in p.runs:
            r.font.size, r.font.bold = Pt(20), True

for table in d.tables:
    # Preserve the raw model response table as a table, alongside the result tables.
    table.autofit = False
    n = len(table.columns)
    widths = (
        [2900, 3300, 3160]
        if n == 3
        else [2400, 1000, 1000, 1000, 1900, 2060]
        if n == 6
        else [9360 // n] * n
    )
    widths[-1] += 9360 - sum(widths)
    pr = table._tbl.tblPr
    for tag in ["tblW", "tblInd", "tblCellMar"]:
        old = pr.find(qn("w:" + tag))
        if old is not None:
            pr.remove(old)
    for tag, width in [("tblW", 9360), ("tblInd", 100)]:
        node = OxmlElement("w:" + tag)
        node.set(qn("w:w"), str(width))
        node.set(qn("w:type"), "dxa")
        pr.append(node)
    mar = OxmlElement("w:tblCellMar")
    for side in ["top", "left", "bottom", "right"]:
        node = OxmlElement("w:" + side)
        node.set(qn("w:w"), "100")
        node.set(qn("w:type"), "dxa")
        mar.append(node)
    pr.append(mar)
    for col, width in zip(table._tbl.tblGrid.gridCol_lst, widths):
        col.set(qn("w:w"), str(width))
    for i, row in enumerate(table.rows):
        rp = row._tr.get_or_add_trPr()
        rp.append(OxmlElement("w:cantSplit"))
        if i == 0:
            rp.append(OxmlElement("w:tblHeader"))
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.space_before = Pt(3)
                for r in p.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(9.5)

# Explicit fonts must take precedence over the reference document's theme fonts.
for style in d.styles:
    if style.font.name == "Arial":
        for fonts in style._element.xpath(".//w:rFonts"):
            for attr in list(fonts.attrib):
                if "theme" in attr.lower():
                    del fonts.attrib[attr]

for section in d.sections:
    p = section.footer.paragraphs[0]
    p.alignment = 2
    r = p.add_run()
    f = OxmlElement("w:fldSimple")
    f.set(qn("w:instr"), "PAGE")
    r._r.addnext(f)
    p.style = d.styles["Caption"]
d.save(sys.argv[2])
